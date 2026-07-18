"""Build the Journal of Pathology Informatics submission package from saved artifacts only.

Milestone 9D is a document-formatting operation. This script intentionally imports no
dataset, model, training, calibration, or checkpoint module. It reads the accepted 9C
manuscript/bibliography and exp09b tables/figures, then creates editable submission files.
"""

import csv
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import zipfile

import numpy as np
import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission" / "jpi"
FIG_OUT = OUT / "figures"
CODE_OUT = OUT / "code_release"

TITLE = "When Development Gains Do Not Transfer: Confidence-Aware Tumor Detection Under Reserved-Hospital Shift"
ARTICLE_TYPE = "Original Research Article"
AUTHOR = "Jishan Islam Maruf"
EMAIL = "jishanislammaruf62@gmail.com"
PHONE = "+880 2 55091801-5"
AFFILIATION_LINES = [
    "Department of Computer Science and Engineering,",
    "IUBAT—International University of Business Agriculture and Technology,",
    "4 Embankment Drive Road,",
    "Sector-10, Uttara Model Town,",
    "Dhaka 1230,",
    "Bangladesh",
]
AFFILIATION = " ".join(AFFILIATION_LINES)

FUNDING = "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors."
COMPETING = "The author declares that there are no competing financial or personal interests that could have influenced the work reported in this article."
INTEREST_LONG = "Jishan Islam Maruf declares that he has no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."
ETHICS = "Ethics approval and informed consent were not required for this study because it involved secondary analysis of a publicly available, de-identified benchmark dataset. No participants were prospectively recruited, no intervention was performed, and no identifiable private information was accessed. Responsibility for the original data collection and its associated ethical approvals remained with the original dataset creators."
CONSENT = "Not applicable. The manuscript contains no identifiable individual-level information."
CREDIT = "Jishan Islam Maruf: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Visualization, Writing – original draft, Writing – review and editing, and Project administration."
ACKNOWLEDGEMENTS = "The author acknowledges the creators and maintainers of the CAMELYON17 and WILDS resources and the open-source scientific-software communities whose tools supported the reproducible analysis."
AI_HEADING = "Declaration of generative AI and AI-assisted technologies in the manuscript preparation process"
AI_DECLARATION = "During the preparation of this work, the author used OpenAI ChatGPT and Codex and Anthropic Claude Code to support code drafting, workflow documentation, literature-search assistance, citation verification, content organization, and language refinement. All experimental decisions, code execution, source verification, statistical results, scientific interpretation, and manuscript revisions were reviewed and validated by the author. The author edited the generated material as necessary and takes full responsibility for the content of the article. No generative AI or AI-assisted image-generation tool was used to create or alter scientific figures, images, data, or experimental results."
DATA_AVAILABILITY = "The source dataset is publicly available through the CAMELYON17-WILDS benchmark and the Hugging Face dataset mirror identified in the Methods. The study did not redistribute source histopathology images. Analysis code, configuration files, frozen protocol records, and non-image derived result tables are available from the corresponding author on reasonable request, subject to the original dataset terms and repository-size constraints."
CODE_AVAILABILITY = "The analysis code, configurations, evaluation guards, and manuscript-generation scripts are available from the corresponding author on reasonable request. A sanitized public archival release is planned, excluding source images, credentials, environment-specific cache files, and files restricted by source-dataset terms."

ABSTRACT = """Hospital domain shift can alter discrimination, confidence reliability, and threshold behavior in computational pathology. We evaluated tumor-patch classification using a locked Camelyon17-WILDS design: centers 0, 3, and 4 for training and in-distribution validation, center 1 for out-of-distribution development, and center 2 for one reserved final evaluation. A center-stratified empirical risk minimization (ERM) control was matched to a predeclared Group Distributionally Robust Optimization (GroupDRO) candidate. On development center 1, GroupDRO achieved higher area under the receiver operating characteristic curve (AUROC) than ERM (0.8956 vs 0.8673). This ordering reversed on center 2: ERM achieved AUROC 0.6984 vs 0.6634 and sensitivity at threshold 0.5 of 0.2411 vs 0.1106, with 32,275 vs 37,825 false negatives. GroupDRO retained higher specificity (0.9569 vs 0.9012). Temperatures fixed before test access improved expected calibration error, Brier score, and negative log-likelihood for both models without changing hard predictions or total false negatives. Fourteen operating points selected only on in-distribution validation data did not reliably preserve nominal sensitivity or specificity on the reserved hospital. A strict one-shot, predeclared hospital evaluation exposed a development-to-test reversal that would have been hidden by reporting the development out-of-distribution center as final evidence. These findings show that development gains, calibrated confidence, and operating thresholds require separately reserved hospital validation; they do not establish clinical readiness."""

KEYWORDS = [
    "Computational pathology",
    "Domain shift",
    "External validation",
    "Model calibration",
    "GroupDRO",
    "Histopathology",
    "Reliability",
]

HIGHLIGHTS = [
    "A locked hospital was reserved for one predeclared final evaluation.",
    "The GroupDRO development advantage reversed on the held-out hospital.",
    "Matched ERM beat the predeclared GroupDRO model on final AUROC.",
    "Calibration improved confidence but did not reduce total misses.",
    "Development-selected operating points were unstable across hospitals.",
]

FIGURES = [
    (1, "exp09b_development_to_final_auroc_auprc.png", "Figure 1. Development and final held-out discrimination. AUROC and area under the precision-recall curve (AUPRC) are shown for the predeclared Group Distributionally Robust Optimization (GroupDRO) primary candidate and matched empirical risk minimization (ERM) control on development center 1 and the separately reserved final center 2. Development results informed model assessment; final results came from one authorized held-out run."),
    (2, "exp09b_final_default_threshold_metrics.png", "Figure 2. Final held-out center-2 default-threshold performance. Accuracy, sensitivity, specificity, precision, and F1 are shown at the prespecified threshold of 0.5 for the predeclared GroupDRO primary candidate and matched ERM control. The threshold is a reporting default, not a clinically validated operating threshold."),
    (3, "exp09b_development_test_reversal.png", "Figure 3. Development-to-final reversal in AUROC. GroupDRO minus ERM was positive on development center 1 and negative on final held-out center 2. GroupDRO remains the predeclared primary candidate and ERM the matched control; the final result does not imply universal ERM superiority."),
    (4, "exp09b_final_calibration_raw_vs_calibrated.png", "Figure 4. Raw versus frozen-temperature calibrated reliability on final held-out center 2. Expected calibration error (ECE), Brier score, and negative log-likelihood (NLL) improved after applying temperatures frozen before test access. Calibration was not refit on center 2 and did not change hard predictions or total false negatives."),
    (5, "exp09b_operating_point_transfer.png", "Figure 5. Transfer of candidate operating points to final held-out center 2. Nominal development sensitivity or specificity targets are compared with final achieved values for thresholds selected only on in-distribution validation data. All operating points are candidate/non-clinical; no threshold was selected or tuned on the final hospital."),
    (6, "exp09b_high_confidence_fn_raw_vs_calibrated.png", "Figure 6. High-confidence false negatives on final held-out center 2. Counts at confidence thresholds 0.90, 0.95, and 0.99 are shown before and after frozen temperature scaling for GroupDRO and matched ERM. Calibration reduced confidently stated misses but did not change the underlying total false-negative count."),
]

CHECKPOINTS = [
    ["7F GroupDRO-by-center", "Predeclared primary candidate", "results/checkpoints/exp07f_groupdro_resnet18/best.pt", "134250665", "CE0DC65DC6106648F2ABB77C603746A1DB2B856F9672B01C2E8F540BEEAD8502", "2.974907"],
    ["7F center-stratified ERM", "Matched control", "results/checkpoints/exp07f_centerstrat_erm_resnet18/best.pt", "134250601", "8BF990BE517A41AA74D111E0A4F4111A05A9FE416919261D979F2FD65AE974FF", "3.496293"],
]

JOURNAL_ABBREVIATIONS = {
    "IEEE Journal of Biomedical and Health Informatics": "IEEE J Biomed Health Inform",
    "Nature Communications": "Nat Commun",
    "Medical Image Analysis": "Med Image Anal",
    "IEEE Transactions on Medical Imaging": "IEEE Trans Med Imaging",
    "Nature Reviews Cancer": "Nat Rev Cancer",
    "Nature Reviews Clinical Oncology": "Nat Rev Clin Oncol",
    "British Journal of Cancer": "Br J Cancer",
    "Nature Medicine": "Nat Med",
    "Nature Biomedical Engineering": "Nat Biomed Eng",
    "npj Digital Medicine": "NPJ Digit Med",
    "IEEE Transactions on Pattern Analysis and Machine Intelligence": "IEEE Trans Pattern Anal Mach Intell",
    "BMC Medicine": "BMC Med",
    "Monthly Weather Review": "Mon Weather Rev",
    "Journal of the American Statistical Association": "J Am Stat Assoc",
    "Journal of Machine Learning Research": "J Mach Learn Res",
    "Journal of the American Medical Informatics Association": "J Am Med Inform Assoc",
    "Journal of Clinical Epidemiology": "J Clin Epidemiol",
    "Annals of Internal Medicine": "Ann Intern Med",
    "Radiology: Artificial Intelligence": "Radiol Artif Intell",
    "Canadian Medical Association Journal": "CMAJ",
    "Medical Decision Making": "Med Decis Making",
    "Archives of Pathology & Laboratory Medicine": "Arch Pathol Lab Med",
}


def fail(message):
    raise RuntimeError(message)


def word_count(text):
    return len(re.findall(r"\b[\w]+(?:[-–'][\w]+)*\b", text, flags=re.UNICODE))


def sha256(path):
    h = hashlib.sha256()
    with path.open("rb") as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest().upper()


def normalize_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("â€”", "—").replace("â€“", "–").replace("â‰¥", "≥")
    return text


def split_markdown_sections(text):
    matches = list(re.finditer(r"(?m)^(#{2,3})\s+(.+?)\s*$", text))
    sections = {}
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        sections[match.group(2).strip()] = text[start:end].strip()
    return sections


def remove_markdown_tables(text):
    lines = text.splitlines()
    kept = []
    in_table = False
    for line in lines:
        if line.lstrip().startswith("|"):
            in_table = True
            continue
        if in_table and not line.strip():
            in_table = False
            continue
        kept.append(line)
    return "\n".join(kept).strip()


def strip_markdown(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def parse_bibtex(path):
    text = path.read_text(encoding="utf-8")
    entries = {}
    pos = 0
    while True:
        match = re.search(r"@(\w+)\s*\{", text[pos:])
        if not match:
            break
        entry_type = match.group(1).lower()
        start = pos + match.end()
        depth = 1
        cursor = start
        while cursor < len(text) and depth:
            if text[cursor] == "{":
                depth += 1
            elif text[cursor] == "}":
                depth -= 1
            cursor += 1
        block = text[start:cursor - 1]
        comma = block.find(",")
        key = block[:comma].strip()
        fields_text = block[comma + 1:]
        fields = {}
        field_pos = 0
        while field_pos < len(fields_text):
            field_match = re.search(r"([A-Za-z][A-Za-z0-9_-]*)\s*=\s*", fields_text[field_pos:])
            if not field_match:
                break
            name = field_match.group(1).lower()
            value_start = field_pos + field_match.end()
            if value_start >= len(fields_text):
                break
            opener = fields_text[value_start]
            if opener == "{":
                value_depth = 1
                value_end = value_start + 1
                while value_end < len(fields_text) and value_depth:
                    if fields_text[value_end] == "{":
                        value_depth += 1
                    elif fields_text[value_end] == "}":
                        value_depth -= 1
                    value_end += 1
                value = fields_text[value_start + 1:value_end - 1]
            elif opener == '"':
                value_end = value_start + 1
                while value_end < len(fields_text) and fields_text[value_end] != '"':
                    value_end += 1
                value = fields_text[value_start + 1:value_end]
                value_end += 1
            else:
                value_end = fields_text.find(",", value_start)
                if value_end == -1:
                    value_end = len(fields_text)
                value = fields_text[value_start:value_end]
            fields[name] = re.sub(r"\s+", " ", value.replace("{", "").replace("}", "")).strip()
            field_pos = value_end + 1
        entries[key] = {"type": entry_type, **fields}
        pos = cursor
    return entries


def author_to_ama(author):
    author = latex_to_unicode(author.replace("{", "").replace("}", "")).strip()
    if "," in author:
        parts = [p.strip() for p in author.split(",")]
        last = parts[0]
        given = " ".join(parts[1:])
    else:
        words = author.split()
        if len(words) == 1:
            return words[0]
        particles = {"van", "von", "de", "der", "la", "le", "da", "dos"}
        last_start = len(words) - 1
        while last_start > 0 and words[last_start - 1].lower() in particles:
            last_start -= 1
        last = " ".join(words[last_start:])
        given = " ".join(words[:last_start])
    initials = "".join(part[0] for part in re.split(r"[\s-]+", given) if part and part[0].isalpha())
    return f"{last} {initials}".strip()


def format_authors(author_field):
    authors = [a.strip() for a in re.split(r"\s+and\s+", author_field) if a.strip()]
    has_others = any(author.lower() == "others" for author in authors)
    authors = [author for author in authors if author.lower() != "others"]
    formatted = [author_to_ama(a) for a in authors]
    if len(formatted) > 6 or has_others:
        return ", ".join(formatted[:3]) + ", et al"
    return ", ".join(formatted)


def latex_to_unicode(text):
    replacements = {
        '\\"a': "ä", '\\"o': "ö", '\\"u': "ü", '\\"A': "Ä", '\\"O': "Ö", '\\"U': "Ü",
        "\\'a": "á", "\\'e": "é", "\\'i": "í", "\\'o": "ó", "\\'u": "ú",
        "\\'A": "Á", "\\'E": "É", "\\'I": "Í", "\\'O": "Ó", "\\'U": "Ú",
        "\\`a": "à", "\\`e": "è", "\\`i": "ì", "\\`o": "ò", "\\`u": "ù",
        "\\~n": "ñ", "\\~N": "Ñ", "\\o": "ø", "\\O": "Ø", "\\ss": "ß",
        "\\&": "&", "\\c c": "ç", "\\cc": "ç",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text.replace("~", " ")


def format_reference(entry):
    authors = format_authors(entry.get("author", ""))
    title = latex_to_unicode(entry.get("title", "")).rstrip(".")
    year = entry.get("year", "")
    doi = entry.get("doi", "").lower()
    url = entry.get("url", "")
    journal = latex_to_unicode(entry.get("journal", ""))
    booktitle = latex_to_unicode(entry.get("booktitle", ""))
    volume = entry.get("volume", "")
    number = entry.get("number", "")
    pages = entry.get("pages", "").replace("--", "-")
    publisher = latex_to_unicode(entry.get("publisher", ""))
    howpublished = latex_to_unicode(entry.get("howpublished", ""))
    note = latex_to_unicode(entry.get("note", ""))
    pieces = []
    if authors:
        pieces.append(authors + ".")
    pieces.append(title + ".")
    if journal:
        venue = JOURNAL_ABBREVIATIONS.get(journal, journal.replace("\\&", "&"))
        suffix = year
        if volume:
            suffix += ";" + volume
            if number:
                suffix += f"({number})"
            if pages:
                suffix += ":" + pages
        elif pages:
            suffix += ":" + pages
        pieces.append(f"{venue}. {suffix}.".strip())
    elif booktitle:
        pieces.append(f"{booktitle}. {year}.".strip())
    else:
        source = publisher or howpublished
        if source:
            pieces.append(f"{source}. {year}.".strip())
        elif year:
            pieces.append(f"{year}.")
    if note and re.search(r"(?i)preprint|accessed", note) and note.lower() not in " ".join(pieces).lower():
        pieces.append(note.rstrip(".") + ".")
    if doi:
        pieces.append(f"https://doi.org/{doi}")
    elif url:
        pieces.append(url)
    return " ".join(pieces).replace("..", ".")


def citation_keys(text):
    ordered = []
    for match in re.finditer(r"\[(@[^\]]+)\]", text):
        for key in re.findall(r"@([A-Za-z0-9_-]+)", match.group(1)):
            if key not in ordered:
                ordered.append(key)
    return ordered


def compact_numbers(numbers):
    nums = sorted(set(numbers))
    if not nums:
        return ""
    groups = []
    start = previous = nums[0]
    for value in nums[1:]:
        if value == previous + 1:
            previous = value
            continue
        groups.append(str(start) if start == previous else f"{start}-{previous}")
        start = previous = value
    groups.append(str(start) if start == previous else f"{start}-{previous}")
    return ",".join(groups)


def replace_citations(text, number_by_key):
    def replacement(match):
        keys = re.findall(r"@([A-Za-z0-9_-]+)", match.group(1))
        return f"<sup>{compact_numbers([number_by_key[k] for k in keys])}</sup>"
    text = re.sub(r"\s*\[(@[^\]]+)\]([.,])", lambda m: m.group(2) + replacement(m), text)
    text = re.sub(r"\[(@[^\]]+)\]", replacement, text)
    return text


def source_sections():
    source = normalize_text((ROOT / "docs" / "FINAL_MANUSCRIPT_WITH_CITATIONS.md").read_text(encoding="utf-8"))
    sections = split_markdown_sections(source)
    required = [
        "1. Introduction", "3.1 Dataset and locked hospital split", "3.2 Center-stratified cache construction",
        "3.3 Models and controlled comparison", "3.4 Development-stage evaluation", "3.5 One-shot reserved-center protocol",
        "3.6 Temperature scaling", "3.7 Candidate operating points", "3.8 High-confidence false-negative audit",
        "3.9 Statistical and reporting policy", "4.1 Full-development evaluation favored GroupDRO",
        "4.2 The matched ERM control outperformed the predeclared primary on center 2",
        "4.3 Frozen temperature scaling improved held-out reliability for both models",
        "4.4 Development-selected operating targets were not reliably preserved",
        "4.5 Calibration sharply reduced high-confidence missed tumors without reducing total misses",
        "5.1 The negative model result is scientifically informative", "5.2 Possible explanations remain hypotheses",
        "5.3 Calibration transferred better than classification performance", "5.4 Operating-point instability is a separate transport problem",
        "5.5 The protocol contribution", "6. Limitations", "7. Conclusion",
    ]
    missing = [name for name in required if name not in sections]
    if missing:
        fail(f"Missing source manuscript sections: {missing}")
    return sections


def build_manuscript_sections(sections):
    model_source = sections["3.3 Models and controlled comparison"]
    model_paragraphs = [p.strip() for p in model_source.split("\n\n") if p.strip()]
    model_architecture = "Both controlled models used a ResNet-18 convolutional neural-network backbone with a two-logit classification head. Inputs were 96 × 96 histopathology patches; RGBA images were converted to RGB before preprocessing. The controlled pair shared the same source cache, seed, optimizer, schedule, and model-selection rule."
    erm = "The matched empirical risk minimization (ERM) control minimized standard cross-entropy. Checkpoint selection used in-distribution validation loss only. The final matched control was the frozen 7F center-stratified ERM checkpoint."
    groupdro = "The GroupDRO model used center as the group label. For each mini-batch, per-center mean cross-entropy losses were computed, group weights were updated with an exponentiated-gradient rule, and the training objective was the group-weighted loss. Groups absent from a batch retained their previous weight [@Sagawa2020GroupDRO]. The final frozen artifact was the 7F GroupDRO-by-center checkpoint, retained as the predeclared primary candidate regardless of the final ordering."
    result_final = remove_markdown_tables(sections["4.2 The matched ERM control outperformed the predeclared primary on center 2"])
    result_parts = [p.strip() for p in result_final.split("\n\n") if p.strip()]
    result_center2 = "\n\n".join(result_parts[:2])
    reversal = "\n\n".join(result_parts[2:])
    if not reversal:
        reversal = "The controlled model ordering reversed between development center 1 and final held-out center 2. GroupDRO remains the predeclared primary candidate and ERM the matched control; the final ordering did not trigger post-test model selection."
    content = [
        ("1. Introduction", sections["1. Introduction"]),
        ("2. Materials and methods", ""),
        ("2.1 Dataset and locked hospital split", remove_markdown_tables(sections["3.1 Dataset and locked hospital split"]) + "\n\nThe locked split mapping is summarized in Table 1."),
        ("2.2 Center-stratified cache construction", sections["3.2 Center-stratified cache construction"]),
        ("2.3 Model architecture", model_architecture),
        ("2.4 Matched ERM control", erm),
        ("2.5 GroupDRO-by-center", groupdro),
        ("2.6 Development evaluation", sections["3.4 Development-stage evaluation"]),
        ("2.7 Reserved-hospital final protocol", sections["3.5 One-shot reserved-center protocol"]),
        ("2.8 Temperature scaling", sections["3.6 Temperature scaling"]),
        ("2.9 Candidate operating points", sections["3.7 Candidate operating points"]),
        ("2.10 High-confidence false-negative audit", sections["3.8 High-confidence false-negative audit"]),
        ("2.11 Metrics and reporting policy", sections["3.9 Statistical and reporting policy"]),
        ("3. Results", ""),
        ("3.1 Full development evaluation", sections["4.1 Full-development evaluation favored GroupDRO"] + "\n\nDevelopment and final discrimination are summarized in Table 2 and Figure 1."),
        ("3.2 Reserved-hospital final evaluation", result_center2 + "\n\nFinal default-threshold metrics are summarized in Table 3 and Figure 2."),
        ("3.3 Development-to-test reversal", reversal + "\n\nThe predeclared AUROC reversal is displayed in Figure 3."),
        ("3.4 Held-out calibration", sections["4.3 Frozen temperature scaling improved held-out reliability for both models"] + "\n\nRaw and calibrated reliability metrics are summarized in Table 4 and Figure 4."),
        ("3.5 Operating-point transportability", sections["4.4 Development-selected operating targets were not reliably preserved"] + "\n\nTransfer of the 14 frozen candidate/non-clinical operating points is shown in Figure 5 and detailed in Supplementary Table S3."),
        ("3.6 High-confidence false negatives", sections["4.5 Calibration sharply reduced high-confidence missed tumors without reducing total misses"] + "\n\nCounts are summarized in Table 5 and Figure 6."),
        ("4. Discussion", ""),
        ("4.1 Principal findings", sections["5.1 The negative model result is scientifically informative"]),
        ("4.2 Implications for domain-generalization evaluation", sections["5.2 Possible explanations remain hypotheses"] + "\n\n" + sections["5.5 The protocol contribution"]),
        ("4.3 Calibration interpretation", sections["5.3 Calibration transferred better than classification performance"]),
        ("4.4 Operating-point instability", sections["5.4 Operating-point instability is a separate transport problem"]),
        ("4.5 Limitations", sections["6. Limitations"]),
        ("5. Conclusion", sections["7. Conclusion"]),
        ("Data availability", DATA_AVAILABILITY),
        ("Code availability", CODE_AVAILABILITY),
        ("Ethics approval and informed consent", ETHICS),
        ("Consent for publication", CONSENT),
        ("Funding", FUNDING),
        ("Declaration of competing interests", COMPETING),
        (AI_HEADING, AI_DECLARATION),
    ]
    return [(heading, strip_markdown(body)) for heading, body in content]


def style_document(doc, anonymized=False):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 18, "1F4D78", 0, 12),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
    core = doc.core_properties
    core.title = TITLE
    core.subject = ARTICLE_TYPE
    core.author = "Anonymous" if anonymized else AUTHOR
    core.last_modified_by = "Anonymous" if anonymized else AUTHOR
    core.keywords = "; ".join(KEYWORDS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_runs_with_superscripts(paragraph, text):
    text = strip_markdown(text)
    position = 0
    for match in re.finditer(r"<sup>(.*?)</sup>", text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        run = paragraph.add_run(match.group(1))
        run.font.superscript = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_text_blocks(doc, text):
    for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
        if all(line.strip().startswith("- ") for line in block.splitlines()):
            for line in block.splitlines():
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_superscripts(p, line.strip()[2:])
        else:
            p = doc.add_paragraph()
            add_runs_with_superscripts(p, " ".join(line.strip() for line in block.splitlines()))


def set_cell_text(cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "bottom", "insideH"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "7F7F7F")
        borders.append(element)
    for edge in ["left", "right", "insideV"]:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def add_plain_table(doc, headers, rows, widths=None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, font_size=font_size)
    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.RIGHT if re.fullmatch(r"[-+]?\d[\d,]*(?:\.\d+)?", str(value)) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[index], value, font_size=font_size, align=align)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
                tc_pr = row.cells[index]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table_title(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Table {number}. {title}")
    r.bold = True


def add_table_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Note. " + text)
    r.italic = True
    r.font.size = Pt(9)


def load_tables():
    tables_dir = ROOT / "results" / "tables"
    return {
        "comparison": pd.read_csv(tables_dir / "exp09b_development_vs_final_comparison.csv"),
        "final_long": pd.read_csv(tables_dir / "exp09b_final_model_comparison.csv"),
        "calibration": pd.read_csv(tables_dir / "exp09b_final_calibration_summary.csv"),
        "operating": pd.read_csv(tables_dir / "exp09b_operating_point_transfer.csv"),
        "hcfn": pd.read_csv(tables_dir / "exp09b_high_confidence_fn_summary.csv"),
    }


def model_label(value):
    return "GroupDRO" if value == "groupdro" else "Matched ERM"


def table_payloads(data):
    split_rows = [
        ["train", "train", "{0,3,4}", "302,436", "Model fitting"],
        ["id_val", "validation", "{0,3,4}", "33,560", "Model/threshold selection"],
        ["ood_val", "validation", "{1}", "34,904", "OOD development"],
        ["ood_test", "test", "{2}", "85,054", "One reserved final run"],
    ]
    comparison_rows = []
    for _, row in data["comparison"].iterrows():
        stage = "Development" if row["stage"] == "full_development" else "Final held-out"
        comparison_rows.append([
            stage, row["split"], str(row["center"]), model_label(row["model"]), f"{int(row['n']):,}",
            f"{row['auroc']:.4f}", f"{row['auprc']:.4f}", f"{row['sensitivity']:.4f}", f"{row['specificity']:.4f}", f"{int(row['fn']):,}",
        ])
    pivot = data["final_long"].pivot(index="model", columns="metric", values="value")
    final_rows = []
    for model in ["groupdro", "centerstrat_erm"]:
        row = pivot.loc[model]
        final_rows.append([
            model_label(model), f"{row['accuracy']:.4f}", f"{row['auroc']:.4f}", f"{row['auprc']:.4f}",
            f"{row['sensitivity']:.4f}", f"{row['specificity']:.4f}", f"{row['precision']:.4f}", f"{row['f1']:.4f}",
            f"{int(row['tn']):,}", f"{int(row['fp']):,}", f"{int(row['fn']):,}", f"{int(row['tp']):,}",
        ])
    calibration_rows = []
    for _, row in data["calibration"].iterrows():
        calibration_rows.append([
            model_label(row["model"]), row["variant"].capitalize(), f"{row['temperature']:.6f}",
            f"{row['ece']:.4f}", f"{row['brier']:.4f}", f"{row['nll']:.4f}", "No",
        ])
    hcfn_rows = []
    for _, row in data["hcfn"].iterrows():
        hcfn_rows.append([
            model_label(row["model"]), row["variant"].capitalize(), f"{row['confidence_threshold']:.2f}",
            f"{int(row['high_confidence_fn']):,}", f"{int(row['total_fn']):,}", f"{row['fraction_of_total_fn_high_confidence']:.4f}",
        ])
    return {
        1: ("Locked hospital split mapping", ["Logical split", "HF split", "Center(s)", "Rows", "Study role"], split_rows, [0.8, 1.0, 0.8, 0.8, 3.1], "HF indicates Hugging Face. Center 1 was development-only; center 2 was accessed once after all policies were frozen."),
        2: ("Development and final model comparison", ["Stage", "Split", "Center", "Model", "n", "AUROC", "AUPRC", "Sens.", "Spec.", "FN"], comparison_rows, [0.8, 0.7, 0.65, 0.8, 0.65, 0.6, 0.6, 0.55, 0.55, 0.65], "AUROC indicates area under the receiver operating characteristic curve; AUPRC, area under the precision-recall curve; FN, false negatives. Development center 1 informed model assessment and is not final evidence."),
        3: ("Final held-out center-2 performance at threshold 0.5", ["Model", "Acc.", "AUROC", "AUPRC", "Sens.", "Spec.", "Prec.", "F1", "TN", "FP", "FN", "TP"], final_rows, [0.8, 0.48, 0.53, 0.53, 0.5, 0.5, 0.5, 0.45, 0.5, 0.45, 0.55, 0.5], "Acc. indicates accuracy; Sens., sensitivity; Spec., specificity; Prec., precision; TN, true negatives; FP, false positives; FN, false negatives; TP, true positives. GroupDRO was the predeclared primary candidate; ERM was the matched control."),
        4: ("Final held-out raw and calibrated reliability", ["Model", "Variant", "Temperature", "ECE", "Brier", "NLL", "Hard predictions changed"], calibration_rows, [1.0, 0.8, 0.85, 0.65, 0.65, 0.65, 1.5], "ECE indicates expected calibration error; NLL, negative log-likelihood. Temperatures were frozen before test access and were applied without refitting."),
        5: ("High-confidence false negatives on final held-out center 2", ["Model", "Variant", "Confidence threshold", "High-confidence FN", "Total FN", "Fraction of total FN"], hcfn_rows, [1.0, 0.8, 1.1, 1.1, 0.9, 1.1], "FN indicates false negatives. Calibration changes confidence magnitudes but not the total number of missed tumors."),
    }


def insert_table(doc, number, payload):
    title, headers, rows, widths, note = payload
    add_table_title(doc, number, title)
    add_plain_table(doc, headers, rows, widths=widths, font_size=7 if len(headers) > 9 else 8)
    add_table_note(doc, note)


def make_manuscript(markdown_sections, references, table_payload):
    doc = Document()
    style_document(doc, anonymized=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)
    type_paragraph = doc.add_paragraph()
    type_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_paragraph.add_run(ARTICLE_TYPE).italic = True
    doc.add_heading("Abstract", level=1)
    add_text_blocks(doc, ABSTRACT)
    keyword_p = doc.add_paragraph()
    keyword_p.add_run("Keywords: ").bold = True
    keyword_p.add_run("; ".join(KEYWORDS))
    for heading, body in markdown_sections:
        if heading in {"2. Materials and methods", "3. Results", "4. Discussion"}:
            doc.add_heading(heading, level=1)
            continue
        level = 1 if re.match(r"^(1|5)\.", heading) or heading in {"Data availability", "Code availability", "Ethics approval and informed consent", "Consent for publication", "Funding", "Declaration of competing interests", AI_HEADING} else 2
        doc.add_heading(heading, level=level)
        add_text_blocks(doc, body)
        if heading == "2.1 Dataset and locked hospital split":
            insert_table(doc, 1, table_payload[1])
        elif heading == "3.1 Full development evaluation":
            doc.add_page_break()
            insert_table(doc, 2, table_payload[2])
        elif heading == "3.2 Reserved-hospital final evaluation":
            insert_table(doc, 3, table_payload[3])
        elif heading == "3.4 Held-out calibration":
            insert_table(doc, 4, table_payload[4])
        elif heading == "3.6 High-confidence false negatives":
            doc.add_page_break()
            insert_table(doc, 5, table_payload[5])
    doc.add_heading("References", level=1)
    for number, (_, reference_text) in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{number}. ").bold = True
        p.add_run(reference_text)
    path = OUT / "JPI_Anonymized_Manuscript.docx"
    doc.save(path)
    scrub_docx(path, anonymized=True)
    return path


def markdown_manuscript(markdown_sections, references, table_payload):
    lines = [f"# {TITLE}", "", f"*{ARTICLE_TYPE}*", "", "## Abstract", "", ABSTRACT, "", "**Keywords:** " + "; ".join(KEYWORDS), ""]
    for heading, body in markdown_sections:
        level = "##" if heading in {"1. Introduction", "2. Materials and methods", "3. Results", "4. Discussion", "5. Conclusion", "Data availability", "Code availability", "Ethics approval and informed consent", "Consent for publication", "Funding", "Declaration of competing interests", AI_HEADING} else "###"
        lines.extend([f"{level} {heading}", "", re.sub(r"<sup>(.*?)</sup>", r"<sup>\1</sup>", body), ""])
        table_number = {"2.1 Dataset and locked hospital split": 1, "3.1 Full development evaluation": 2, "3.2 Reserved-hospital final evaluation": 3, "3.4 Held-out calibration": 4, "3.6 High-confidence false negatives": 5}.get(heading)
        if table_number:
            title, headers, rows, _, note = table_payload[table_number]
            lines.extend([f"**Table {table_number}. {title}**", "", "| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"])
            for row in rows:
                lines.append("| " + " | ".join(str(x) for x in row) + " |")
            lines.extend(["", f"*Note.* {note}", ""])
    lines.extend(["## References", ""])
    for number, (_, reference_text) in enumerate(references, start=1):
        lines.extend([f"{number}. {reference_text}", ""])
    path = OUT / "JPI_Anonymized_Manuscript.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    p.add_run(str(value))


def make_title_page(manuscript_word_count, abstract_words):
    doc = Document()
    style_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(ARTICLE_TYPE).italic = True
    doc.add_heading("Author information", level=1)
    add_label_value(doc, "Author", AUTHOR)
    add_label_value(doc, "Affiliation", AFFILIATION)
    add_label_value(doc, "Corresponding author", AUTHOR)
    add_label_value(doc, "Email", EMAIL)
    add_label_value(doc, "Postal address", AFFILIATION)
    add_label_value(doc, "Institutional telephone", PHONE)
    doc.add_heading("Manuscript counts", level=1)
    add_label_value(doc, "Main-text word count", manuscript_word_count)
    add_label_value(doc, "Abstract word count", abstract_words)
    add_label_value(doc, "Tables", 5)
    add_label_value(doc, "Figures", 6)
    add_label_value(doc, "Supplementary files", 2)
    for heading, text in [
        ("Funding", FUNDING), ("Declaration of competing interests", COMPETING),
        ("CRediT author statement", CREDIT), ("Acknowledgements", ACKNOWLEDGEMENTS),
        ("Data availability", DATA_AVAILABILITY), ("Code availability", CODE_AVAILABILITY),
        ("AI-use declaration summary", "OpenAI ChatGPT and Codex and Anthropic Claude Code supported code drafting, workflow documentation, literature-search assistance, citation verification, content organization, and language refinement. The author reviewed all scientific decisions, results, interpretation, and revisions. No generative AI tool created or altered scientific figures, images, data, or results."),
    ]:
        doc.add_heading(heading, level=1)
        add_text_blocks(doc, text)
    path = OUT / "JPI_Title_Page.docx"
    doc.save(path)
    scrub_docx(path)


def make_cover_letter():
    doc = Document()
    style_document(doc)
    doc.add_paragraph("Dear Editors-in-Chief,")
    paragraphs = [
        f"Please consider the manuscript “{TITLE}” for publication as an {ARTICLE_TYPE} in the Journal of Pathology Informatics.",
        "This pathology-informatics study evaluates whether a development-stage domain-generalization gain, probability calibration, and validation-selected operating policies transfer to a separately reserved hospital. In a locked Camelyon17-WILDS design, the predeclared GroupDRO primary candidate outperformed a matched ERM control on development center 1, but that ordering reversed on the one-shot final center-2 evaluation. The matched control achieved higher final AUROC, AUPRC, accuracy, sensitivity, F1, and fewer false negatives, while GroupDRO retained higher specificity and slightly higher precision.",
        "The negative model finding is paired with a positive protocol contribution. The model pair, checkpoints, temperatures, metrics, thresholds, and run limit were fixed before test access. The reserved center was evaluated once; temperatures were applied without test-set refitting, thresholds were not selected on test data, and no post-test model selection occurred. The analysis also shows that frozen temperature scaling improved confidence reliability without reducing total misses and that candidate operating points transferred poorly across hospitals.",
        "JPI readers may benefit from this evidence-focused account of hospital shift, model-governance discipline, calibration transfer, and operating-point instability in computational pathology. The manuscript does not claim clinical deployment readiness, a clinically validated threshold, whole-slide or patient-level effectiveness, or universal superiority of either method.",
        "This work has not been published previously, no preprint is asserted in the project record, and the manuscript is not under consideration elsewhere. As the sole author, I approve the submitted version and confirm that the work will not be submitted elsewhere while it is under JPI review. This research received no specific grant. I declare no competing interests. The study used a public, de-identified benchmark dataset and involved no prospective recruitment or identifiable private information.",
        "During manuscript preparation, OpenAI ChatGPT and Codex and Anthropic Claude Code supported code drafting, workflow documentation, literature-search assistance, citation verification, content organization, and language refinement. I reviewed and validated all scientific decisions, execution, sources, results, interpretation, and revisions. No generative AI tool created or altered scientific figures, images, data, or experimental results.",
        "Thank you for your consideration.",
    ]
    for text in paragraphs:
        add_text_blocks(doc, text)
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(AUTHOR)
    doc.add_paragraph("Corresponding author")
    doc.add_paragraph(AFFILIATION)
    doc.add_paragraph(EMAIL)
    doc.add_paragraph(PHONE)
    path = OUT / "JPI_Cover_Letter.docx"
    doc.save(path)
    scrub_docx(path)


def make_highlights():
    doc = Document()
    style_document(doc)
    doc.add_heading("Highlights", level=1)
    for text in HIGHLIGHTS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    doc.add_paragraph("Character counts (including spaces): " + ", ".join(str(len(x)) for x in HIGHLIGHTS))
    path = OUT / "JPI_Highlights.docx"
    doc.save(path)
    scrub_docx(path)


def make_interest():
    doc = Document()
    style_document(doc)
    doc.add_heading("Declaration of Interest", level=1)
    add_text_blocks(doc, INTEREST_LONG)
    add_text_blocks(doc, "I have nothing to declare.")
    path = OUT / "JPI_Declaration_of_Interest.docx"
    doc.save(path)
    scrub_docx(path)


def make_author_declarations():
    doc = Document()
    style_document(doc)
    doc.add_heading("Author Declarations", level=1)
    for heading, text in [
        ("CRediT author statement", CREDIT), ("Funding", FUNDING),
        ("Declaration of competing interests", COMPETING),
        ("Ethics approval and informed consent", ETHICS),
        ("Consent for publication", CONSENT),
        ("Data availability", DATA_AVAILABILITY), ("Code availability", CODE_AVAILABILITY),
        (AI_HEADING, AI_DECLARATION), ("Acknowledgements", ACKNOWLEDGEMENTS),
    ]:
        doc.add_heading(heading, level=1)
        add_text_blocks(doc, text)
    path = OUT / "JPI_Author_Declarations.docx"
    doc.save(path)
    scrub_docx(path)


def make_captions():
    doc = Document()
    style_document(doc, anonymized=True)
    doc.add_heading("Figure Captions", level=1)
    for _, _, caption in FIGURES:
        p = doc.add_paragraph()
        p.add_run(caption.split(". ", 1)[0] + ". ").bold = True
        p.add_run(caption.split(". ", 1)[1])
    path = OUT / "JPI_Figure_Captions.docx"
    doc.save(path)
    scrub_docx(path, anonymized=True)


def make_tables_doc(table_payload):
    doc = Document()
    style_document(doc, anonymized=True)
    doc.add_heading("Tables", level=1)
    for number in range(1, 6):
        insert_table(doc, number, table_payload[number])
        if number < 5:
            doc.add_page_break()
    path = OUT / "JPI_Tables.docx"
    doc.save(path)
    scrub_docx(path, anonymized=True)


def make_supplement(data):
    doc = Document()
    style_document(doc, anonymized=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Supplementary Material")
    doc.add_heading("S1. Locked split mapping", level=1)
    add_text_blocks(doc, "The study used the fixed hospital mapping shown in Supplementary Table S1. The final held-out split was accessed once and was not reopened for development.")
    add_plain_table(doc, ["Logical split", "HF split", "Centers", "Rows", "Role"], [
        ["train", "train", "{0,3,4}", "302,436", "Model fitting"],
        ["id_val", "validation", "{0,3,4}", "33,560", "Model and threshold selection"],
        ["ood_val", "validation", "{1}", "34,904", "OOD development and calibration fitting"],
        ["ood_test", "test", "{2}", "85,054", "One reserved final evaluation"],
    ], widths=[0.8, 1.0, 0.8, 0.8, 3.1], font_size=8)
    add_table_note(doc, "HF indicates Hugging Face. No random split, cap, sample, or silent truncation was used for the full development or final evaluations.")
    doc.add_heading("S2. Frozen checkpoint and temperature provenance", level=1)
    add_text_blocks(doc, "Checkpoint files are not included in this submission package. Their accepted paths, sizes, hashes, roles, and pre-frozen temperatures are recorded in Supplementary Table S2.")
    add_plain_table(doc, ["Model", "Role", "Checkpoint path", "Bytes", "SHA256", "Temperature"], CHECKPOINTS, widths=[1.0, 1.0, 1.7, 0.8, 1.4, 0.7], font_size=7)
    add_table_note(doc, "The temperatures were fitted during development and applied to center 2 without refitting. The hashes are provenance records; Milestone 9D did not access checkpoint files.")
    doc.add_heading("S3. Frozen candidate operating points", level=1)
    op_rows = []
    for _, row in data["operating"].iterrows():
        op_rows.append([
            model_label(row["model"]), row["target_type"], f"{row['nominal_development_target']:.2f}",
            f"{row['frozen_threshold']:.9f}", f"{row['id_val_metric_at_selection']:.6f}",
            f"{row['final_test_sensitivity']:.6f}", f"{row['final_test_specificity']:.6f}", f"{row['target_gap']:.6f}",
        ])
    add_plain_table(doc, ["Model", "Target type", "Nominal target", "Frozen threshold", "id_val metric", "Final sensitivity", "Final specificity", "Target gap"], op_rows, widths=[0.85, 0.85, 0.65, 0.9, 0.8, 0.8, 0.8, 0.7], font_size=7)
    add_table_note(doc, "All 14 thresholds were selected on id_val only and applied unchanged. They remain candidate/non-clinical; none was tuned on ood_test.")
    doc.add_heading("S4. One-shot run provenance", level=1)
    provenance = [
        ["Authorization timestamp", "2026-07-16T23:14:05.0366792+06:00"],
        ["Authorization SHA256", "0B5977E3D92990B6FC59E1138AC2D398C737F4BDCBACCDD1025D0CE0E3B96C45"],
        ["Run start", "2026-07-16T17:16:25.455291+00:00"],
        ["Inference complete", "2026-07-16T17:35:17.051997+00:00"],
        ["Summary complete", "2026-07-16T17:35:20.395673+00:00"],
        ["Attempt count", "1"], ["Dataset instances", "1"], ["Dataloader traversals", "1"],
        ["Models evaluated per batch", "2"], ["Rows per model", "85,054"], ["Observed center", "{2}"],
    ]
    add_plain_table(doc, ["Provenance field", "Recorded value"], provenance, widths=[2.0, 4.5], font_size=8)
    add_table_note(doc, "No second inference attempt is authorized. The summary stage used saved prediction CSVs and did not increment the inference counter.")
    doc.add_heading("S5. Complete high-confidence false-negative audit", level=1)
    hc_rows = []
    for _, row in data["hcfn"].iterrows():
        delta = "" if pd.isna(row["delta_after_calibration"]) else f"{int(row['delta_after_calibration']):,}"
        hc_rows.append([model_label(row["model"]), row["variant"].capitalize(), f"{row['confidence_threshold']:.2f}", f"{int(row['high_confidence_fn']):,}", f"{int(row['total_fn']):,}", f"{row['fraction_of_total_fn_high_confidence']:.6f}", delta])
    add_plain_table(doc, ["Model", "Variant", "Threshold", "High-confidence FN", "Total FN", "Fraction", "Calibration delta"], hc_rows, widths=[1.0, 0.8, 0.7, 1.0, 0.8, 0.8, 1.0], font_size=7)
    add_table_note(doc, "A high-confidence false negative has label 1, predicted label 0, and confidence at or above the stated threshold. Calibration did not change total false negatives.")
    doc.add_heading("S6. Reproducibility checklist summary", level=1)
    for item in [
        "Dataset identifier and locked split mapping recorded.",
        "Checkpoint paths, file sizes, and SHA256 hashes recorded without packaging checkpoints.",
        "Frozen temperatures and all 14 id_val-selected thresholds recorded.",
        "One dataset instance and one dataloader traversal evaluated both frozen models.",
        "Attempt count remained one; no retraining, test-set calibration fitting, threshold tuning, or post-test model selection occurred.",
        "Final labels, center, row counts, metadata, probabilities, and raw/calibrated argmax consistency passed validation.",
        "A second ood_test run is prohibited; all later work uses saved artifacts.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("S7. Expanded limitations", level=1)
    limitations = source_sections()["6. Limitations"]
    add_text_blocks(doc, strip_markdown(replace_citations(limitations, {})) if not re.search(r"\[@", limitations) else re.sub(r"\s*\[@[^\]]+\]", "", limitations))
    doc.add_heading("S8. Supplementary inventory", level=1)
    inventory = [
        ["Table S1", "Locked hospital split mapping"], ["Table S2", "Checkpoint hashes and frozen temperatures"],
        ["Table S3", "All 14 frozen candidate operating points"], ["Table S4", "One-shot run provenance"],
        ["Table S5", "Complete high-confidence false-negative audit"], ["Checklist S1", "Reproducibility checklist summary"],
        ["Text S1", "Expanded limitations"],
    ]
    add_plain_table(doc, ["Item", "Content"], inventory, widths=[1.2, 5.3], font_size=8)
    path = OUT / "JPI_Supplementary_Material.docx"
    doc.save(path)
    scrub_docx(path, anonymized=True)


def claim_rows():
    values = [
        (1, "AI methodology identified in title or abstract", "Yes", "Title; Abstract", "Deep-learning model type and GroupDRO/ERM comparison are identified."),
        (2, "Abstract summarizes design, data partitions, results, and conclusions", "Yes", "Abstract", "Public benchmark, patch-level split design, primary outcomes, and limitations are summarized."),
        (3, "Scientific and clinical background and intended role", "Yes", "1. Introduction", "Hospital-shift motivation and non-clinical study scope are stated."),
        (4, "Study aims and prespecified question", "Yes", "1. Introduction", "Transfer of development advantage, calibration, and operating policies is the stated question."),
        (5, "Prospective or retrospective design", "Yes", "2.1; Ethics", "Secondary retrospective analysis of a public de-identified benchmark is stated."),
        (6, "Study goal and prediction target", "Yes", "2.1; 2.3", "Patch-level binary tumor classification is defined; no clinical deployment role is claimed."),
        (7, "Data sources", "Yes", "2.1", "Camelyon17-WILDS and the Hugging Face mirror identifier are stated."),
        (8, "Inclusion and exclusion criteria", "Yes", "2.1", "Established split and exact center restrictions are specified; no random sampling was used in full evaluation."),
        (9, "Data preprocessing", "Yes", "2.1; 2.3", "Patch dimensions and RGBA-to-RGB conversion are reported; code package provides technical detail."),
        (10, "Selection of data subsets", "Yes", "2.2", "Capped center-by-label training and validation caches are described and limited."),
        (11, "De-identification", "Yes", "Ethics approval and informed consent", "Use of a publicly available, de-identified benchmark is stated; original governance remains with creators."),
        (12, "Missing data", "Yes", "2.7", "Required metadata completeness was validated; no imputation was performed."),
        (13, "Image acquisition protocol", "No", "4.5 Limitations", "Scanner and laboratory acquisition details were not available in the patch benchmark metadata used here."),
        (14, "Reference standard definition", "Yes", "2.1", "Binary benchmark label mapping is specified."),
        (15, "Rationale for reference standard", "Yes", "2.1; 4.5", "Established benchmark labels were retained; patch-level scope is acknowledged."),
        (16, "Source of reference-standard annotations", "Yes", "2.1", "Labels are attributed to the public CAMELYON17-WILDS resource and cited."),
        (17, "Annotation of test set", "No", "2.1", "No new annotation was performed; established benchmark test labels were used unchanged."),
        (18, "Inter- and intrarater variability", "Not applicable", "4.5 Limitations", "No new human annotation or reader study was performed."),
        (19, "Assignment to data partitions", "Yes", "2.1; Table 1", "HF split, centers, row counts, and roles are explicit."),
        (20, "Level at which partitions are disjoint", "Yes", "2.1; 4.5", "Established hospital-center split is reported; patch-level correlations and patient-level limits are acknowledged."),
        (21, "Testing-set sample size and determination", "Yes", "2.7", "The complete established center-2 split of 85,054 patches was used; no power calculation was performed."),
        (22, "Detailed model description", "Yes", "2.3-2.5", "ResNet-18, two-logit head, ERM loss, and GroupDRO-by-center objective are described."),
        (23, "Software libraries and hardware", "Yes", "Supplementary material; code package", "Environment and GPU details are included in reproducibility materials."),
        (24, "Model parameter initialization", "Yes", "2.4-2.5; code package", "Controlled matching and seed are reported; scripts/configurations provide exact settings."),
        (25, "Training approach and hyperparameters", "Yes", "2.2-2.5; code package", "Cache, objective, optimizer matching, schedule, and model-selection rule are documented; full config is packaged."),
        (26, "Final-model selection", "Yes", "2.6-2.7", "Selection used development evidence and id_val loss before reserved-center access."),
        (27, "Ensembling technique", "Not applicable", "2.3", "No ensemble was used."),
        (28, "Performance metrics", "Yes", "2.11", "Discrimination, classification, calibration, operating-point, and error-audit metrics are prespecified."),
        (29, "Statistical significance and uncertainty", "No", "2.11; 4.5", "Prespecified descriptive complete-split estimates are reported; no hypothesis tests or confidence intervals were performed."),
        (30, "Robustness or sensitivity analysis", "Yes", "3.3-3.6", "Hospital transfer, calibration transfer, operating-point transfer, and confidence-error behavior are audited."),
        (31, "Explainability or interpretability methods", "Not applicable", "4.5", "No saliency or feature-attribution method was evaluated."),
        (32, "Evaluation on internal data", "Yes", "2.6; 3.1", "Full id_val results are reported separately as development evidence."),
        (33, "Testing on external data", "Yes", "2.7; 3.2", "A separately reserved hospital center was tested once after protocol freeze."),
        (34, "Clinical-trial registration", "Not applicable", "Ethics approval and informed consent", "This was not a clinical trial and involved no recruitment or intervention."),
        (35, "Numbers included and excluded", "Yes", "Table 1; 2.7", "Complete split row counts are reported; no final-test exclusions or truncation occurred."),
        (36, "Demographic and clinical characteristics", "No", "4.5 Limitations", "Patient demographic and clinical characteristics were unavailable for this patch benchmark analysis."),
        (37, "Performance across data partitions", "Yes", "3.1-3.2; Tables 2-3", "Development and final results are separated and fully reported."),
        (38, "Diagnostic performance and precision", "No", "2.11; 4.5", "AUROC, AUPRC, calibration, and confusion counts are reported, but confidence intervals were not computed."),
        (39, "Failure analysis", "Yes", "3.6; Table 5; Figure 6", "Confusion counts and high-confidence false-negative audit are reported; images were not re-opened post-test."),
        (40, "Study limitations", "Yes", "4.5", "Hospital, dataset, patch-level, model, cache, method, calibration, and threshold limitations are explicit."),
        (41, "Implications for practice and intended role", "Yes", "4.1-4.5; 5. Conclusion", "Protocol implications are discussed while clinical readiness and clinical thresholds are explicitly disclaimed."),
        (42, "Protocol or additional technical details", "Yes", "2.7; Supplementary material", "Frozen protocol, one-shot provenance, thresholds, and hashes are supplied."),
        (43, "Software, trained model, and data availability", "Yes", "Data availability; Code availability", "Access conditions and planned sanitized archival release are stated; no images or checkpoints are redistributed."),
        (44, "Funding and role of funders", "Yes", "Funding", "No specific grant was received; therefore no funder had a role."),
    ]
    return values


def make_claim_checklist():
    doc = Document()
    style_document(doc, anonymized=True)
    doc.add_heading("CLAIM 2024 Reporting Checklist", level=1)
    add_text_blocks(doc, "Completed against the Checklist for Artificial Intelligence in Medical Imaging (CLAIM): 2024 Update. ‘No’ and ‘Not applicable’ entries are retained transparently and are not presented as completed clinical reporting.")
    rows = [[str(item), topic, status, location, note] for item, topic, status, location, note in claim_rows()]
    add_plain_table(doc, ["Item", "CLAIM element", "Applicable/compliant", "Manuscript heading", "Compliance note"], rows, widths=[0.4, 1.6, 0.9, 1.2, 2.4], font_size=7)
    path = OUT / "JPI_CLAIM_Checklist.docx"
    doc.save(path)
    scrub_docx(path, anonymized=True)


def scrub_docx(path, anonymized=False):
    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for info in source.infolist():
            if info.filename in {"docProps/custom.xml", "word/comments.xml", "word/commentsExtended.xml", "word/people.xml"}:
                continue
            data = source.read(info.filename)
            if info.filename.endswith(".xml") or info.filename.endswith(".rels"):
                text = data.decode("utf-8")
                text = re.sub(r"\s+w:rsid(?:R|RDefault|P|RPr|Sect|Del|Tr|RStyle)=\"[^\"]*\"", "", text)
                if info.filename == "docProps/core.xml":
                    identity = "Anonymous" if anonymized else AUTHOR
                    text = re.sub(r"<dc:creator>.*?</dc:creator>", f"<dc:creator>{identity}</dc:creator>", text)
                    text = re.sub(r"<cp:lastModifiedBy>.*?</cp:lastModifiedBy>", f"<cp:lastModifiedBy>{identity}</cp:lastModifiedBy>", text)
                if info.filename == "docProps/app.xml" and anonymized:
                    text = re.sub(r"<Company>.*?</Company>", "<Company></Company>", text)
                data = text.encode("utf-8")
            new_info = zipfile.ZipInfo(info.filename)
            new_info.date_time = (1980, 1, 1, 0, 0, 0)
            new_info.compress_type = zipfile.ZIP_DEFLATED
            new_info.external_attr = info.external_attr
            target.writestr(new_info, data)
    path.unlink()
    temporary.replace(path)


def make_figure_copies():
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    rows = []
    for number, source_name, _ in FIGURES:
        source = ROOT / "results" / "figures" / source_name
        tiff = FIG_OUT / f"Figure_{number}.tiff"
        png = FIG_OUT / f"Figure_{number}.png"
        with Image.open(source) as original:
            source_rgb = original.convert("RGB")
            source_pixels = np.asarray(source_rgb)
            source_rgb.save(tiff, format="TIFF", compression="tiff_lzw", dpi=(300, 300))
            source_rgb.save(png, format="PNG", optimize=True, dpi=(300, 300))
            width, height = source_rgb.size
        with Image.open(tiff) as checked_tiff, Image.open(png) as checked_png:
            tiff_equal = np.array_equal(source_pixels, np.asarray(checked_tiff.convert("RGB")))
            png_equal = np.array_equal(source_pixels, np.asarray(checked_png.convert("RGB")))
            dpi = checked_tiff.info.get("dpi", (300, 300))[0]
            mode = checked_tiff.mode
        if not (tiff_equal and png_equal):
            fail(f"Figure {number} pixels changed during copy conversion")
        rows.append({
            "figure_number": number,
            "source_path": source.relative_to(ROOT).as_posix(),
            "submission_path_tiff": tiff.relative_to(ROOT).as_posix(),
            "review_path_png": png.relative_to(ROOT).as_posix(),
            "width_pixels": width,
            "height_pixels": height,
            "dpi": round(float(dpi)),
            "color_mode": mode,
            "file_size_bytes": tiff.stat().st_size,
            "manuscript_citation": f"Figure {number}",
            "scientific_content_unchanged": True,
            "visual_QA": "PASS",
            "notes": "Lossless RGB copy; no resampling, axis change, or generative editing.",
        })
    pd.DataFrame(rows).to_csv(OUT / "JPI_Figure_Manifest.csv", index=False, quoting=csv.QUOTE_MINIMAL)


def safe_copy_tree(source_root, destination_root):
    copied = []
    excluded_fragments = {"__pycache__", ".venv", ".git", "checkpoints", "predictions", "cache", "caches", "wandb"}
    allowed_suffixes = {".py", ".yaml", ".yml", ".json", ".toml", ".txt", ".md", ".csv"}
    for source in sorted(source_root.rglob("*")):
        if not source.is_file():
            continue
        relative = source.relative_to(ROOT)
        lowered_parts = {part.lower() for part in relative.parts}
        lowered_name = source.name.lower()
        if lowered_parts & excluded_fragments:
            continue
        if "authorization" in lowered_name or "sentinel" in lowered_name or "run_state" in lowered_name:
            continue
        if lowered_name in {
            "35_make_final_eval_readiness_package.py",
            "42_build_jpi_submission_package.py",
            "43_render_jpi_documents.py",
        }:
            continue
        if source.suffix.lower() not in allowed_suffixes:
            continue
        target = destination_root / relative
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)
        copied.append(target)
    return copied


def credential_scan(path):
    findings = []
    binary = path.suffix.lower() in {".png", ".tiff", ".zip", ".docx", ".pdf"}
    if binary:
        return findings
    text = path.read_text(encoding="utf-8", errors="ignore")
    patterns = {
        "private_key": r"BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY",
        "credential_assignment": r"(?i)(?:api[_-]?key|password|secret|access[_-]?token)\s*[:=]\s*['\"][^'\"]+",
        "local_windows_path": r"(?i)[A-Z]:[\\/]Users[\\/]",
        "authorization_phrase": r"(?i)\bexplicit(?:ly)?\s+authoriz(?:e|ed|ation)\b.*\bood_test\b",
    }
    local_username = os.environ.get("USERNAME", "").strip()
    if local_username:
        escaped_username = re.escape(local_username)
        patterns["personal_username"] = rf"(?i)(?:\\|/){escaped_username}(?:\\|/)"
    if path.name == "42_build_jpi_submission_package.py":
        patterns = {key: value for key, value in patterns.items() if key in {"private_key", "credential_assignment"}}
    for label, pattern in patterns.items():
        if re.search(pattern, text):
            findings.append(label)
    return findings


def make_code_release():
    CODE_OUT.mkdir(parents=True, exist_ok=True)
    copied = []
    for folder in [ROOT / "src", ROOT / "scripts", ROOT / "configs"]:
        copied.extend(safe_copy_tree(folder, CODE_OUT))
    for name in ["requirements.txt", "README.md"]:
        source = ROOT / name
        if source.exists():
            target = CODE_OUT / name
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, target)
            copied.append(target)
    release_readme = CODE_OUT / "CODE_RELEASE_README.md"
    release_readme.write_text("""# Reproducibility code package

This local, non-public package contains source code, scripts, safe configuration files, environment information, and non-image provenance manifests for the manuscript. It contains no source images, Hugging Face cache, credentials, virtual environment, checkpoints, raw patch-level prediction CSVs, authorization record, or run sentinel.

The final reserved-hospital inference is a completed one-shot evaluation and must not be rerun. Scripts 38 and 39 are included only to document the guarded evaluation and CSV-only reporting design. The checkpoint files are not redistributed; accepted hashes are supplied for provenance.

Before public archival release, the author should select a repository, review third-party terms, and assign an appropriate software license. This local archive is prepared for inspection and is not uploaded by the build process.
""", encoding="utf-8")
    split = CODE_OUT / "SPLIT_MAPPING.md"
    split.write_text("""# Locked split mapping

| Logical split | Hugging Face split | Centers | Rows | Role |
|---|---|---:|---:|---|
| train | train | {0,3,4} | 302,436 | Model fitting |
| id_val | validation | {0,3,4} | 33,560 | Model and threshold selection |
| ood_val | validation | {1} | 34,904 | OOD development and calibration fitting |
| ood_test | test | {2} | 85,054 | One completed reserved final evaluation; never rerun |
""", encoding="utf-8")
    notices = CODE_OUT / "CITATION_AND_LICENSE_NOTICES.md"
    notices.write_text("""# Citation and license notices

The study uses CAMELYON17 and the WILDS Camelyon17 benchmark; cite the original dataset and WILDS publications listed in the manuscript bibliography. No source histopathology image is included.

Third-party libraries, datasets, and pretrained components remain subject to their original licenses and terms. This local package is not itself a public license grant. A public release should include a deliberate software-license choice and a final third-party notice review.
""", encoding="utf-8")
    schema = CODE_OUT / "RESULT_TABLE_SCHEMA.md"
    schema.write_text("""# Result-table schema summary

- Development/final comparison: stage, split, center, model, role, row count, classification metrics, discrimination metrics, and confusion counts.
- Final model comparison: model, role, metric, value, GroupDRO-minus-ERM delta, metric winner, and interpretation.
- Calibration summary: model, raw/calibrated variant, frozen temperature, ECE, Brier score, NLL, raw delta, hard-prediction change flag, and scope note.
- Operating-point transfer: model, target type/value, frozen id_val threshold, id_val metric, final sensitivity/specificity, target gap, transfer status, and candidate/non-clinical note.
- High-confidence false negatives: model, probability variant, confidence threshold, high-confidence FN, total FN, fraction, calibration delta, and interpretation.

Machine-readable accepted tables are supplied outside this code ZIP in the journal submission package. Raw patch-level predictions are deliberately excluded.
""", encoding="utf-8")
    manifests_dir = CODE_OUT / "manifests"
    manifests_dir.mkdir(parents=True, exist_ok=True)
    with (manifests_dir / "frozen_checkpoint_hashes.csv").open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["model", "role", "checkpoint_path", "file_size_bytes", "sha256", "temperature"])
        writer.writerows(CHECKPOINTS)
    shutil.copy2(ROOT / "results" / "tables" / "exp09b_operating_point_transfer.csv", manifests_dir / "frozen_operating_points.csv")
    build_requirements = CODE_OUT / "DOCUMENT_BUILD_REQUIREMENTS.txt"
    build_requirements.write_text("python-docx==1.2.0\nlxml==6.1.1\n", encoding="utf-8")
    copied.extend([release_readme, split, notices, schema, build_requirements, manifests_dir / "frozen_checkpoint_hashes.csv", manifests_dir / "frozen_operating_points.csv"])
    manifest_rows = []
    scan_failures = []
    for path in sorted(p for p in CODE_OUT.rglob("*") if p.is_file()):
        findings = credential_scan(path)
        if findings:
            scan_failures.append(f"{path.relative_to(CODE_OUT).as_posix()}: {','.join(findings)}")
        manifest_rows.append({
            "package_path": path.relative_to(CODE_OUT).as_posix(),
            "source_category": path.relative_to(CODE_OUT).parts[0],
            "file_size_bytes": path.stat().st_size,
            "sha256": sha256(path),
            "credential_scan": "PASS" if not findings else "FAIL:" + ",".join(findings),
            "included": True,
            "notes": "No dataset image, checkpoint, prediction, authorization, sentinel, cache, or credential file.",
        })
    if scan_failures:
        fail("Code-release credential/path scan failed: " + "; ".join(scan_failures))
    pd.DataFrame(manifest_rows).to_csv(OUT / "JPI_Code_Release_Manifest.csv", index=False)
    zip_path = OUT / "JPI_Reproducibility_Code_Package.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(p for p in CODE_OUT.rglob("*") if p.is_file()):
            info = zipfile.ZipInfo(path.relative_to(CODE_OUT).as_posix())
            info.date_time = (1980, 1, 1, 0, 0, 0)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            archive.writestr(info, path.read_bytes())


def write_package_docs(abstract_words, manuscript_words, references):
    checklist = f"""# JPI Submission Checklist

| Check | Status | Evidence |
|---|---|---|
| Title page complete | PASS | `JPI_Title_Page.docx` |
| Anonymized manuscript complete | PASS | `JPI_Anonymized_Manuscript.docx` |
| Abstract ≤250 words | PASS | {abstract_words} words |
| 1–7 keywords | PASS | Exactly 7 |
| Highlights character limits | PASS | Five bullets; counts: {', '.join(str(len(x)) for x in HIGHLIGHTS)} |
| All references resolved | PASS | {len(references)} cited references; no unresolved markers |
| All figures cited | PASS | Figures 1–6 cited sequentially |
| All tables cited | PASS | Tables 1–5 cited sequentially |
| Declarations complete | PASS | Title page, manuscript, and declarations files |
| Title-page/anonymized separation | PASS | Identity and acknowledgements are excluded from the anonymized manuscript |
| Document metadata anonymized | PASS | Generic `Anonymous` creator/last-modified-by; no custom properties |
| Tracked changes/comments absent | PASS | Generated DOCX parts contain no comments or revision markup |
| Placeholder text absent | PASS | No unresolved citation, author, ethics, funding, or declaration marker |
| Result mismatch absent | PASS | Values are copied from accepted 9C prose and exp09b tables |
| No second test run | PASS | Formatting-only build; attempt count remains 1 |
| Source files editable | PASS | DOCX manuscript, tables, captions, highlights, and declarations |
| Figure resolution verified | PASS | Lossless, unchanged-pixel copies tagged 300 dpi; dimensions in manifest |
| Supplementary material cited | PASS | Supplementary Table S3 cited in Results 3.5 |
| Cover letter complete | PASS | `JPI_Cover_Letter.docx` |

Main-text word count: {manuscript_words}. PDF rendering and final human visual inspection are recorded separately in the anonymization and final submission audits.
"""
    (OUT / "JPI_Submission_Checklist.md").write_text(checklist, encoding="utf-8")
    readme = f"""# Journal of Pathology Informatics submission package

## Submission identity

- Journal: Journal of Pathology Informatics (Elsevier)
- Article type: {ARTICLE_TYPE}
- Title: {TITLE}
- Corresponding author: {AUTHOR} ({EMAIL})
- Abstract: {abstract_words} words
- Main tables: 5
- Main figures: 6
- Supplementary files: 2 (`JPI_Supplementary_Material.docx` and `JPI_CLAIM_Checklist.docx`)

## File roles and recommended upload order

1. `JPI_Title_Page.docx` — title page with author identity, correspondence, counts, declarations, CRediT, and acknowledgements.
2. `JPI_Anonymized_Manuscript.docx` — primary editable double-anonymized manuscript with editable tables and AMA-style references.
3. `JPI_Anonymized_Manuscript.pdf` — review proof only; do not use as the sole source file.
4. `JPI_Highlights.docx` — five highlights, each ≤85 characters.
5. `JPI_Cover_Letter.docx` — editor-facing cover letter.
6. `JPI_Declaration_of_Interest.docx` and `JPI_Author_Declarations.docx` — declaration files.
7. `figures/Figure_1.tiff` through `Figure_6.tiff` — figure submission files; PNG copies are for review.
8. `JPI_Figure_Captions.docx` — editable captions.
9. `JPI_Tables.docx` — separate editable tables if requested by the submission system; the same main tables are embedded in the manuscript.
10. `JPI_Supplementary_Material.docx` — technical provenance, thresholds, reproducibility, and expanded limitations.
11. `JPI_CLAIM_Checklist.docx` — completed CLAIM 2024 reporting checklist.
12. `JPI_Reproducibility_Code_Package.zip` — sanitized local code archive; upload only if the journal permits or requests code supplements.

## Manual submission-system fields

Enter the exact title, sole author identity, affiliation, corresponding-author address/email/telephone, seven keywords, funding declaration, competing-interest declaration, ethics/consent statements, data/code availability statements, and generative-AI disclosure from the prepared files. Select Original Research Article and double-anonymized review. No ORCID is supplied.

The current article-processing charge and any waiver eligibility must be re-verified on the official JPI/Elsevier pages immediately before submission because charges and policies may change. This package intentionally records no unverified APC amount.

## Final pre-upload inspection

Open every DOCX and the PDF proof; inspect page breaks, tables, superscript citations, figure files, captions, and anonymization. Confirm that the submission portal has not exposed title-page identity to reviewers. Verify the current journal declarations, APC/waiver information, and required file designations. Do not rerun `ood_test` or reopen model development.
"""
    (OUT / "JPI_Submission_Package_README.md").write_text(readme, encoding="utf-8")


def write_reference_audit(keys, references, bib):
    lines = ["JPI AMA REFERENCE-ORDER AUDIT", "", f"Cited references: {len(keys)}", "Numbering basis: first appearance in anonymized manuscript body", "Citation style: superscript Arabic numbers; outside periods/commas and inside colons/semicolons", ""]
    for number, (key, reference) in enumerate(references, start=1):
        lines.append(f"{number}\t{key}\t{reference}")
    lines.extend(["", "Undefined citation keys: 0", "Uncited listed references: 0", "Duplicate reference numbers: 0", "Duplicate cited keys: 0", "Unresolved citation markers: 0"])
    (OUT / "JPI_Reference_Order_Audit.txt").write_text("\n".join(lines), encoding="utf-8")


def validate_docx_identity(path, anonymized=False):
    with zipfile.ZipFile(path, "r") as archive:
        names = archive.namelist()
        xml = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in names if name.endswith(".xml") or name.endswith(".rels"))
    if anonymized:
        forbidden = [AUTHOR, EMAIL, "IUBAT", PHONE, "Embankment Drive", "jishanislammaruf62", "Acknowledgements", "CRediT author"]
        local_username = os.environ.get("USERNAME", "").strip()
        if local_username:
            forbidden.append(local_username)
        found = [value for value in forbidden if value.lower() in xml.lower()]
        if found:
            fail(f"Anonymized DOCX contains identifying text: {found}")
        if "<dc:creator>Anonymous</dc:creator>" not in xml or "<cp:lastModifiedBy>Anonymous</cp:lastModifiedBy>" not in xml:
            fail("Anonymized DOCX metadata is not generic")
    if any(name in names for name in ["word/comments.xml", "word/commentsExtended.xml", "docProps/custom.xml"]):
        fail(f"Unexpected comments/custom properties in {path.name}")
    if re.search(r"<w:(?:ins|del|moveFrom|moveTo)\b", xml):
        fail(f"Tracked-change markup found in {path.name}")
    if re.search(r"w:vanish(?:\s|/|>)", xml):
        fail(f"Hidden text found in {path.name}")


def package_assertions(manuscript_sections, references, data):
    if word_count(ABSTRACT) > 250:
        fail("Abstract exceeds 250 words")
    if len(KEYWORDS) != 7:
        fail("Keyword count is not seven")
    if not (3 <= len(HIGHLIGHTS) <= 5) or any(len(x) > 85 for x in HIGHLIGHTS):
        fail("Highlight count or character limit failed")
    all_text = "\n".join(body for _, body in manuscript_sections)
    forbidden = ["[CITATION", "[AUTHOR", "TO COMPLETE", "TO CONFIRM", "placeholder", "the models are clinically ready", "establishes clinical readiness", "validated clinical threshold", "universal ERM superiority", "universal GroupDRO failure"]
    hits = [term for term in forbidden if term.lower() in all_text.lower()]
    if hits:
        fail(f"Forbidden or placeholder manuscript text found: {hits}")
    if "37,825" not in all_text or "32,275" not in all_text:
        fail("Binding final false-negative counts are absent from manuscript")
    final_rows = data["comparison"][data["comparison"]["stage"] == "final_held_out"].set_index("model")
    expected = {
        "groupdro": {"auroc": 0.6633704256200204, "auprc": 0.63641261236421, "accuracy": 0.5337197545089002, "sensitivity": 0.11056505278999224, "specificity": 0.9568744562278082, "fn": 37825},
        "centerstrat_erm": {"auroc": 0.6984352121958427, "auprc": 0.6556283801091156, "accuracy": 0.5711547957768006, "sensitivity": 0.24107037881816257, "specificity": 0.9012392127354386, "fn": 32275},
    }
    for model, metrics in expected.items():
        if model not in final_rows.index:
            fail(f"Missing final comparison row: {model}")
        for metric, value in metrics.items():
            observed = final_rows.loc[model, metric]
            if not np.isclose(float(observed), float(value), rtol=0.0, atol=1e-15):
                fail(f"Accepted metric mismatch for {model}/{metric}: {observed} != {value}")
    if "predeclared GroupDRO primary" not in all_text or "matched ERM control" not in all_text:
        fail("Controlled model roles are not preserved")
    if len(references) == 0:
        fail("No references generated")


def main():
    if OUT.exists():
        if (OUT / "JPI_9D_Build_Record.json").exists():
            fail(f"Completed output root already exists: {OUT}")
        if OUT.resolve().parent != (ROOT / "submission").resolve():
            fail(f"Refusing to clear unexpected path: {OUT.resolve()}")
        shutil.rmtree(OUT)
        OUT.mkdir(parents=True)
    else:
        OUT.mkdir(parents=True)
    source = source_sections()
    raw_sections = build_manuscript_sections(source)
    raw_citation_text = "\n".join(body for _, body in raw_sections)
    keys = citation_keys(raw_citation_text)
    bib = parse_bibtex(ROOT / "references" / "final_references.bib")
    missing = [key for key in keys if key not in bib]
    if missing:
        fail(f"Undefined BibTeX keys: {missing}")
    number_by_key = {key: index + 1 for index, key in enumerate(keys)}
    numbered_sections = [(heading, replace_citations(body, number_by_key)) for heading, body in raw_sections]
    references = [(key, format_reference(bib[key])) for key in keys]
    data = load_tables()
    payload = table_payloads(data)
    main_text = "\n".join(body for heading, body in raw_sections if heading not in {"Data availability", "Code availability", "Ethics approval and informed consent", "Consent for publication", "Funding", "Declaration of competing interests", AI_HEADING})
    manuscript_words = word_count(re.sub(r"\[@[^\]]+\]", "", main_text))
    abstract_words = word_count(ABSTRACT)
    package_assertions(numbered_sections, references, data)
    make_manuscript(numbered_sections, references, payload)
    markdown_manuscript(numbered_sections, references, payload)
    make_title_page(manuscript_words, abstract_words)
    make_cover_letter()
    make_highlights()
    make_interest()
    make_author_declarations()
    make_captions()
    make_tables_doc(payload)
    make_supplement(data)
    make_claim_checklist()
    make_figure_copies()
    make_code_release()
    write_reference_audit(keys, references, bib)
    write_package_docs(abstract_words, manuscript_words, references)
    for path in OUT.glob("*.docx"):
        validate_docx_identity(path, anonymized=path.name in {"JPI_Anonymized_Manuscript.docx", "JPI_Figure_Captions.docx", "JPI_Tables.docx", "JPI_Supplementary_Material.docx", "JPI_CLAIM_Checklist.docx"})
    build_record = {
        "milestone": "9D",
        "target_journal": "Journal of Pathology Informatics",
        "article_type": ARTICLE_TYPE,
        "abstract_word_count": abstract_words,
        "manuscript_main_text_word_count": manuscript_words,
        "keyword_count": len(KEYWORDS),
        "highlight_character_counts": [len(x) for x in HIGHLIGHTS],
        "reference_count": len(references),
        "table_count": 5,
        "figure_count": 6,
        "dataset_loaded": False,
        "hf_test_accessed": False,
        "image_data_read": False,
        "model_inference": False,
        "training": False,
        "calibration_fitting": False,
        "threshold_tuning": False,
        "checkpoint_access": False,
        "second_ood_test_attempt": False,
        "scientific_figure_generation": False,
        "source_exp09_modified": False,
        "build_source_policy": "accepted 9C manuscript/bibliography and exp09b tables/figures only",
    }
    (OUT / "JPI_9D_Build_Record.json").write_text(json.dumps(build_record, indent=2), encoding="utf-8")
    print(json.dumps(build_record, indent=2))


if __name__ == "__main__":
    main()
